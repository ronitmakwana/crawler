from encodings import utf_8
from bs4 import BeautifulSoup
import requests
import pdfkit
import os
import uuid
from docx import Document
from docx2pdf import convert

document = Document()



r = requests.get('https://idrw.org/')
os.mkdir('images')
htmlcontent = r.content
# print(htmlcontent)

soup = BeautifulSoup(htmlcontent,"html.parser")
# print(soup.get_text().replace(' ',''))




title = soup.find('div',class_='art-layout-cell art-content')

link  = title.find_all('a')

# print(link)

head = title.find_all('h2')
for h in head:
    title = h.text
    # bold = '\33[1m'
    print(f'title:{title}')
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text(title)
    # file = 'imafd.docx'
    # with open(file,"a") as f:
    #     f.write(f'title:{title}'+'\n'+'\n')

    
               
    for link in h.find_all('a'):
        li = link.get('href')
            # print(li)
        links = requests.get(li)
        htmllink = links.content
        soup2 = BeautifulSoup(htmllink,"html.parser")
            # con = soup2.find('div',class_='art-layout-cell art-content')

        con = soup2.find('div',class_='art-content-layout')
        img = con.find('figure',class_='aligncenter size-large is-resized')
        src = img.select_next('img')
        
        # k = []
        # newsrc = images = src[0]
        # image = newsrc.attrs['src']
        # im = requests.get(image).content
        
        # filename = 'images/result_'+str(uuid.uuid4())+'.jpg'
        # with open(filename,"wb+") as ik:
        #     ik.write(im)
        #     r.add_picture(ik)
            # document.save('demo3.docx')

       
            
           
        # s = con.find_all_next('p')
        # span = con.find('span')
        # t = span.text
        # print(f'Date: {t}')
        # r.add_text(t)
        #     # f.write(f'Date: {t}'+'\n')
            
        # for k in s:
        #     content = k.get_text()
        #     print(content)
        #     r.add_text(content)
        #     document.save('image2.docx')
os.rmdir('images')
convert('image.docx','cov.pdf') 
#                 f.write(content+'\n')
                
            

    


#     # pdf.add_page()
#     # pdf.set_font("Arial",size=12)
#     # pdf.cell(200,10, txt=titles,ln=1,align='C') 
#     # pdf.output("title.pdf")
# print('save')
# # def save_pdf(htmls, file_name):
    
# htmls = file
# file_name = "iaf.pdf"
# options = {

#     'page-size': 'Letter',

#     'margin-top': '0.075in',

#     'margin-right': '0.75in',

#     'margin-bottom': '0.75in',

#     'margin-left': '0.75in',

#     'encoding': "UTF-8",

#     'custom-header': [

#     ('Accept-Encoding', 'gzip')

# ],

#     'cookie': [

#         ('cookie-name1', 'cookie-value1'),

#         ('cookie-name2', 'cookie-value2'),

#         ],

#         'outline-depth': 10,

#         }

# pdfkit.from_file(htmls,file_name, options=options) 
print("pdf saved")
# os.remove(file)

    # print(h.text.split())

# for l in link:  

#     print(l.get('href'))
# a = title.h2.a.text
# print(a)

# for index,item in enumerate(title):
#     content = item.find_all('a')
#     for link in content:
#         print(link.get('href'))
   

    
    # for j in content:
    #     c = j.find('p')
    #     print(f'content:{c}')
# a = title.find('a')
# print(a)
# print(soup.prettify)
# print(a)
# for i in a:
#     ti = i.find('p')
#     print(ti)



# for i in title:
#     i = i.find('title')
#     print(i)




