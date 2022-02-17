from bs4 import BeautifulSoup
import requests
import uuid
from docx import Document
from docx.shared import Pt,Inches,Mm,RGBColor
from docx2pdf import convert
from subprocess import  Popen
import os
import shutil
from PIL import Image
import docx
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH 
import time

document = Document()
#p = document.add_paragraph('Defence_world.net',style='BodyText')

p = document.add_heading('Defence_world.net',0)
p.alignment = 1
p = document.add_paragraph()
r = p.add_run()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(18)
# font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
section = document.sections[0]
section.page_height = Mm(350)
section.page_width = Mm(400)
table = document.add_table(rows=1, cols=0)
row = table.rows[0]

#docx = 'idrw/defnew'+str(uuid.uuid4())+'.docx'
#docx = 'idrw/result_{}.docx'.format(str(time.time()))
al = True
while al:
    fn = input("enter filename ::> ")
   
    docx = 'idrw/'+str(fn)+'.docx'
    fp = 'idrw/'+str(fn)+'.pdf'
    if os.path.isfile(fp):
        print('file already exists','\n')
    if not os.path.isfile(fp):         
       
   
        docx = 'idrw/'+str(fn)+'.docx'
        fp = 'idrw/'+str(fn)+'.pdf'
        al = False
        print('----')
imageee = '/home/bisag/Downloads/web-scrapping-main/image/'
if os.path.exists(imageee):
    shutil.rmtree('image')
    os.mkdir('image')

else:
    os.mkdir('image')
    print('no folder')
imm = '/home/bisag/Downloads/web-scrapping-main/im/'
if os.path.exists(imm):
    shutil.rmtree('im')
    os.mkdir('im')

else:
    os.mkdir('im')
    print('no folder')
kgg = '/home/bisag/Downloads/web-scrapping-main/kg/'
if os.path.exists(kgg):
    shutil.rmtree('kg')
    os.mkdir('kg')
else:
    os.mkdir('kg')
    print('no folder')

r = requests.get("https://www.defenseworld.net/")

html_content = r.content

soup = BeautifulSoup(html_content,"html.parser")
a  = soup.find("div",class_='col-sm-16 col-xs-16')
link = a.find('a')
href = link.get('href')
print(href)
r2 = requests.get('https://www.defenseworld.net/'+str(href))
html_content2 = r2.content
soup2 =BeautifulSoup(html_content2,"html.parser")
title = soup2.find('div',class_="col-md-12 col-sm-16 col-xs-16")
title2 = title.find_all('h4',class_="mt-0")
for l in title2:
    l2 = l.find('a')
    link3 = l2.get('href')
    # print('https://www.defenseworld.net/'+str(link3))
    r3 = requests.get('https://www.defenseworld.net/'+str(link3))
    html_content3 = r3.content
    # print(html_content3)
    soup3 = BeautifulSoup(html_content3,'html.parser')
    # print(soup3)
    titles = soup3.find('h1',class_='mt-0')
    tit = titles.text
    p = document.add_paragraph()
    r = p.add_run()
    r.bold = True
    r.font.size = Pt(26)
    r.underline = True
    
    r.font.color.rgb = RGBColor(133, 33, 23)
    r.add_text('Title : '+tit)
    print(tit)
    date = soup3.find('ul',class_='news-meta')
    dt = date.text
    p = document.add_paragraph()
    r = p.add_run()
    r.bold = True
    r.underline = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(233, 43, 23)
    r.add_text('Date : '+dt)
    print(dt)
    src = soup3.select('img')
    newsrc = images = src[0]
    image = newsrc.attrs['src']
    print(image)
    im = requests.get(image,stream=True).content
    filename = str(imageee)+'/result_'+str(uuid.uuid4())+'.jpg'
    with open(filename,"wb+") as ik:
        ik.write(im)
        # r.add_picture(ik)

        ff = Image.open(ik).convert('P').save(str(imm)+str(uuid.uuid4())+'.png')
        source = imm
        destination = kgg
  
        allfiles = os.listdir(source)
  
        for f in allfiles:
            shutil.move(source + f, destination + f)
            p = document.add_paragraph()
            r = p.add_run()
            r.add_picture(destination + f,width=Inches(12),height=Inches(4))
            p = document.add_paragraph()
            r = p.add_run()
                # document.save('page2.docx')
            print('-------')
    con = soup3.find('div',class_='media-content mt-20 mb-20')
    cont = con.text
    p = document.add_paragraph()
    r = p.add_run()
    r.font.size = Pt(20)
    r.add_text(cont)
    p = document.add_paragraph()
    r = p.add_run()
    p = document.add_paragraph()
    r = p.add_run()
   
    # print(cont)

    document.save(docx)
    # p = con.find_all('p')
    # for pc in p:
    #     print(pc.get_text())
    # time = date.find('li')
    # print(time)
    # da  = time.find('span',class_='ion-android-data icon')
    # print(da.text)
    # for ti in titles:
    #     print(ti)


LIBRE_OFFICE = '/usr/lib/libreoffice/program/soffice'

def convert_to_pdf(input_docx, out_folder):
    p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    print([LIBRE_OFFICE, '--convert-to', 'pdf', input_docx])
    p.communicate()


sample_doc = docx
out_folder = 'idrw'
convert_to_pdf(sample_doc, out_folder)



shutil.rmtree(imageee)
shutil.rmtree(imm)
shutil.rmtree(kgg)
os.remove(docx)
print("pdf saved")
